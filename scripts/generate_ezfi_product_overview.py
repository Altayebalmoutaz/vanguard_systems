#!/usr/bin/env python3
"""Generate the ezFi Executive Product Overview (.docx)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

# ── Brand ────────────────────────────────────────────────────────────────────
BLUE = RGBColor(0x1E, 0x73, 0xE8)
BLUE_DARK = RGBColor(0x0F, 0x4C, 0xA8)
BLUE_MID = RGBColor(0x2B, 0x6C, 0xB0)
GRAY = RGBColor(0x5A, 0x6A, 0x7A)
GRAY_LIGHT = RGBColor(0xF2, 0xF5, 0xF8)
GRAY_LINE = RGBColor(0xD0, 0xD7, 0xDE)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x66, 0xCC, 0x33)
ACCENT_BG = "E8F1FC"
LIGHT_BG = "F2F5F8"
CALLOUT_BG = "EEF4FB"
HEADER_BG = "1E73E8"

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "eligibility_dashboard" / "public" / "ezfi-logo.png"
ASSETS = Path("/tmp/ezfi-doc-assets")
OUT = ROOT / "docs" / "product" / "ezFi_Executive_Product_Overview.docx"

FOOTER_TEXT = "Powered by Smile Suites"


# ── Low-level helpers ────────────────────────────────────────────────────────
def set_run_font(run, *, size=11, bold=False, color=BLACK, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="D0D7DE", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def set_table_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def add_page_number(paragraph):
    """Append PAGE field to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar2)


def add_horizontal_line(paragraph, color="1E73E8", sz="12"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(p, before=0, after=6, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


# ── Content builders ─────────────────────────────────────────────────────────
def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12 if level == 1 else 8, after=4)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE_DARK, name="Calibri")
        add_horizontal_line(p, "1E73E8", "14")
    elif level == 2:
        set_run_font(run, size=12, bold=True, color=BLUE, name="Calibri")
    else:
        set_run_font(run, size=11, bold=True, color=BLUE_MID, name="Calibri")
    return p


def add_body(doc, text, *, size=10.5, bold=False, color=BLACK, after=4, align="left"):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=after, line=1.08)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, *, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=1, line=1.05)
    p.paragraph_format.left_indent = Inches(0.22 + level * 0.18)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10, bold=True, color=BLACK)
        r2 = p.add_run(text)
        set_run_font(r2, size=10, color=GRAY)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10, color=BLACK)
    return p


def add_callout(doc, title: str, body: str, accent=HEADER_BG, *, spacer=True):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_BG)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcBorders"):
            tcPr.remove(child)
    borders = OxmlElement("w:tcBorders")
    for edge, sz, col in (
        ("top", "4", "D0D7DE"),
        ("left", "24", accent),
        ("bottom", "4", "D0D7DE"),
        ("right", "4", "D0D7DE"),
    ):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), col)
        borders.append(el)
    tcPr.append(borders)

    cell.width = Inches(6.7)
    p1 = cell.paragraphs[0]
    set_paragraph_spacing(p1, before=4, after=1)
    r1 = p1.add_run(title)
    set_run_font(r1, size=10, bold=True, color=BLUE_DARK)

    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=4)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.5, color=GRAY)
    if spacer:
        sp = doc.add_paragraph()
        set_paragraph_spacing(sp, before=0, after=2)
    return table


def add_kpi_row(doc, items: list[tuple[str, str]]):
    """items: list of (label, value) shown as metric cards."""
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    remove_table_borders(table)
    width = Inches(6.7 / len(items))
    for i, (label, value) in enumerate(items):
        cell = table.cell(0, i)
        cell.width = width
        shade_cell(cell, LIGHT_BG)
        set_cell_borders(cell, "D0D7DE", "4")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p1, before=5, after=1)
        # Support multi-line values with soft breaks
        lines = value.split("\n")
        for li, line in enumerate(lines):
            if li:
                p1.add_run().add_break()
            r1 = p1.add_run(line)
            set_run_font(r1, size=11, bold=True, color=BLUE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, before=0, after=5)
        r2 = p2.add_run(label)
        set_run_font(r2, size=8, color=GRAY)
    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before=0, after=2)
    return table


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    table.autofit = False

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, HEADER_BG)
        set_cell_borders(cell, "0F4CA8", "4")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=2, after=2)
        r = p.add_run(h)
        set_run_font(r, size=9, bold=True, color=WHITE)

    for ri, row in enumerate(rows):
        bg = LIGHT_BG if ri % 2 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            shade_cell(cell, bg)
            set_cell_borders(cell, "D0D7DE", "4")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=2, after=2)
            r = p.add_run(val)
            set_run_font(r, size=9, color=BLACK)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before=0, after=2)
    return table


def add_flow_boxes(doc, steps: list[str], *, vertical=True):
    """Simple vertical/horizontal process flow using table cells."""
    if vertical:
        table = doc.add_table(rows=len(steps) * 2 - 1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed(table)
        remove_table_borders(table)
        for i, step in enumerate(steps):
            cell = table.rows[i * 2].cells[0]
            shade_cell(cell, ACCENT_BG if i % 2 == 0 else LIGHT_BG)
            set_cell_borders(cell, "1E73E8", "8")
            cell.width = Inches(5.2)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=4, after=4)
            r = p.add_run(step)
            set_run_font(r, size=10, bold=True, color=BLUE_DARK)
            if i < len(steps) - 1:
                arrow = table.rows[i * 2 + 1].cells[0]
                arrow.width = Inches(5.2)
                p2 = arrow.paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p2, before=0, after=0)
                r2 = p2.add_run("▼")
                set_run_font(r2, size=10, bold=True, color=BLUE)
    else:
        n = len(steps)
        table = doc.add_table(rows=1, cols=n * 2 - 1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed(table)
        remove_table_borders(table)
        for i, step in enumerate(steps):
            cell = table.rows[0].cells[i * 2]
            shade_cell(cell, ACCENT_BG)
            set_cell_borders(cell, "1E73E8", "8")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=6, after=6)
            r = p.add_run(step)
            set_run_font(r, size=9, bold=True, color=BLUE_DARK)
            if i < n - 1:
                arrow = table.rows[0].cells[i * 2 + 1]
                p2 = arrow.paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run("→")
                set_run_font(r2, size=14, bold=True, color=BLUE)
    doc.add_paragraph()
    return table


def add_architecture_diagram(doc):
    """Platform architecture as nested styled tables."""
    # Top: OpenDental
    t1 = doc.add_table(rows=1, cols=1)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t1.cell(0, 0)
    shade_cell(cell, HEADER_BG)
    set_cell_borders(cell, "0F4CA8", "8")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=8)
    r = p.add_run("OpenDental  ·  System of Record")
    set_run_font(r, size=12, bold=True, color=WHITE)

    arrow = doc.add_paragraph()
    arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(arrow, before=2, after=2)
    ar = arrow.add_run("▼")
    set_run_font(ar, size=14, bold=True, color=BLUE)

    # Middle: AI Platform with modules
    t2 = doc.add_table(rows=3, cols=1)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(t2)
    header = t2.cell(0, 0)
    shade_cell(header, "0F4CA8")
    set_cell_borders(header, "0F4CA8", "8")
    ph = header.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(ph, before=6, after=6)
    rh = ph.add_run("ezFi AI Platform")
    set_run_font(rh, size=12, bold=True, color=WHITE)

    # Current capabilities
    cur = t2.cell(1, 0)
    shade_cell(cur, ACCENT_BG)
    set_cell_borders(cur, "1E73E8", "6")
    pc = cur.paragraphs[0]
    set_paragraph_spacing(pc, before=4, after=2)
    rc = pc.add_run("  Current Capabilities")
    set_run_font(rc, size=10, bold=True, color=BLUE_DARK)
    for label in (
        "✓  Eligibility Verification  (Flagship)",
        "✓  Voice Verification",
        "✓  OpenDental Writeback",
    ):
        px = cur.add_paragraph()
        set_paragraph_spacing(px, before=1, after=1)
        rx = px.add_run(f"     {label}")
        set_run_font(rx, size=10, color=BLACK)

    # Future
    fut = t2.cell(2, 0)
    shade_cell(fut, LIGHT_BG)
    set_cell_borders(fut, "D0D7DE", "6")
    pf = fut.paragraphs[0]
    set_paragraph_spacing(pf, before=4, after=2)
    rf = pf.add_run("  Revenue Cycle Roadmap")
    set_run_font(rf, size=10, bold=True, color=GRAY)
    for label in (
        "○  Coding Agent",
        "○  Prior Authorization",
        "○  ERA / EOB Processing & Reconciliation",
        "○  Denial Management",
    ):
        px = fut.add_paragraph()
        set_paragraph_spacing(px, before=1, after=1)
        rx = px.add_run(f"     {label}")
        set_run_font(rx, size=10, color=GRAY)

    arrow2 = doc.add_paragraph()
    arrow2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(arrow2, before=2, after=2)
    ar2 = arrow2.add_run("▼")
    set_run_font(ar2, size=14, bold=True, color=BLUE)

    # Bottom: Dashboard + Staff
    t3 = doc.add_table(rows=1, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(t3)
    for i, label in enumerate(("ezFi Dashboard", "Staff Review & Approval")):
        cell = t3.cell(0, i)
        shade_cell(cell, HEADER_BG if i == 0 else "2B6CB0")
        set_cell_borders(cell, "0F4CA8", "8")
        cell.width = Inches(3.2)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=8, after=8)
        r = p.add_run(label)
        set_run_font(r, size=11, bold=True, color=WHITE)

    doc.add_paragraph()


def add_layer_diagram(doc):
    layers = [
        ("01", "Intake & Validation", "Patient · Insurance · Carrier validation against OpenDental records"),
        ("02", "Electronic Eligibility", "Stedi 270/271 inquiry with payer routing and response capture"),
        ("03", "Normalization & Interpretation", "Canonical benefits model · coverage percentages · accumulators"),
        ("04", "Business Rules & Confidence", "Payer rules · caching · confidence scoring · exception routing"),
        ("05", "Human Review", "Dashboard review queue · exception queue · approval workflow"),
        ("06", "OpenDental Writeback", "Benefit Notes · InsVerify · Commlog · Benefits Grid · Adjustments"),
    ]
    table = doc.add_table(rows=len(layers), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    for i, (num, title, desc) in enumerate(layers):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        shade_cell(c0, HEADER_BG)
        shade_cell(c1, ACCENT_BG if i % 2 == 0 else LIGHT_BG)
        set_cell_borders(c0, "0F4CA8", "4")
        set_cell_borders(c1, "D0D7DE", "4")
        c0.width = Inches(0.7)
        c1.width = Inches(5.8)
        c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(num)
        set_run_font(r0, size=14, bold=True, color=WHITE)
        p1 = c1.paragraphs[0]
        set_paragraph_spacing(p1, before=4, after=1)
        r1 = p1.add_run(title)
        set_run_font(r1, size=11, bold=True, color=BLUE_DARK)
        p2 = c1.add_paragraph()
        set_paragraph_spacing(p2, before=0, after=4)
        r2 = p2.add_run(desc)
        set_run_font(r2, size=9.5, color=GRAY)
    doc.add_paragraph()


def add_timeline(doc):
    phases = [
        ("CURRENT", "1E73E8", [
            "Eligibility Verification",
            "Voice Verification",
            "OpenDental Writeback",
        ]),
        ("NEXT", "2B6CB0", [
            "Coding Agent",
            "Prior Authorization",
            "Claim Preparation",
        ]),
        ("FUTURE", "5A6A7A", [
            "ERA Processing",
            "EOB Automation",
            "Denial Management",
            "Revenue Analytics",
            "Enterprise DSO Features",
        ]),
    ]
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    for i, (title, color, items) in enumerate(phases):
        h = table.cell(0, i)
        shade_cell(h, color)
        set_cell_borders(h, color, "6")
        h.width = Inches(2.15)
        ph = h.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(ph, before=8, after=8)
        rh = ph.add_run(title)
        set_run_font(rh, size=12, bold=True, color=WHITE)

        b = table.cell(1, i)
        shade_cell(b, LIGHT_BG)
        set_cell_borders(b, "D0D7DE", "6")
        b.width = Inches(2.15)
        first = True
        for item in items:
            if first:
                p = b.paragraphs[0]
                first = False
            else:
                p = b.add_paragraph()
            set_paragraph_spacing(p, before=3, after=3)
            r = p.add_run(f"•  {item}")
            set_run_font(r, size=10, color=BLACK)
        # trailing spacer
        sp = b.add_paragraph()
        set_paragraph_spacing(sp, before=0, after=4)
    doc.add_paragraph()


def add_benefit_grid(doc):
    items = [
        ("Active Coverage", "Coverage Percentages", "Deductibles"),
        ("Remaining Deductibles", "Annual Maximums", "Remaining Maximums"),
        ("Preventive / Basic / Major", "Orthodontic Benefits", "Implant Benefits"),
        ("Waiting Periods", "Frequency Limitations", "Age Limitations"),
        ("Missing Tooth Clauses", "Alternate Benefits", "Downgrade Rules"),
        ("Prior Auth Indicators", "Individual vs Family Accumulators", "Last Service Dates"),
        ("Subscriber Information", "Dependent Information", "Patient Estimates"),
    ]
    table = doc.add_table(rows=len(items), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    for ri, row in enumerate(items):
        for ci, text in enumerate(row):
            cell = table.cell(ri, ci)
            shade_cell(cell, ACCENT_BG if (ri + ci) % 2 == 0 else LIGHT_BG)
            set_cell_borders(cell, "D0D7DE", "4")
            cell.width = Inches(2.15)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=5, after=5)
            r = p.add_run(text)
            set_run_font(r, size=9.5, bold=True, color=BLUE_DARK)
    doc.add_paragraph()


# ── Diagram images (Pillow) ──────────────────────────────────────────────────
def _font(size=18, bold=False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_rounded_rect(draw, xy, radius, fill, outline=None, width=2):
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def make_traditional_flow_image(path: Path):
    w, h = 900, 620
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    font = _font(17, bold=True)
    font_sm = _font(12)
    steps = [
        "Patient Arrival",
        "Insurance Verification Request",
        "Manual Phone Calls to Payer",
        "Benefit Interpretation",
        "Manual OpenDental Updates",
        "Treatment Estimates",
        "Front Desk Communication",
    ]
    box_w, box_h = 440, 48
    x = (w - box_w) // 2
    y = 16
    for i, step in enumerate(steps):
        color = "#1E73E8" if i in (0, 6) else "#E8F1FC"
        text_c = "#FFFFFF" if i in (0, 6) else "#0F4CA8"
        outline = "#0F4CA8"
        draw_rounded_rect(draw, (x, y, x + box_w, y + box_h), 8, color, outline, 2)
        bbox = draw.textbbox((0, 0), step, font=font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        draw.text(((w - tw) // 2, y + (box_h - th) // 2 - 2), step, fill=text_c, font=font)
        if i < len(steps) - 1:
            ay = y + box_h
            draw.line((w // 2, ay + 2, w // 2, ay + 18), fill="#1E73E8", width=3)
            draw.polygon(
                [(w // 2 - 7, ay + 14), (w // 2 + 7, ay + 14), (w // 2, ay + 24)],
                fill="#1E73E8",
            )
            y = ay + 28
        else:
            y += box_h + 8
    note = "High touch  ·  Error-prone  ·  Slow  ·  Inconsistent documentation"
    bbox = draw.textbbox((0, 0), note, font=font_sm)
    tw = bbox[2] - bbox[0]
    draw.text(((w - tw) // 2, h - 28), note, fill="#5A6A7A", font=font_sm)
    img.save(path, "PNG")


def make_voice_flow_image(path: Path):
    w, h = 1200, 220
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    font = _font(15, bold=True)
    steps = [
        "Electronic\nEligibility",
        "Incomplete\nInformation",
        "AI Voice Agent\nContacts Payer",
        "Dashboard\nReview",
        "Approval",
        "OpenDental\nWriteback",
    ]
    box_w, box_h = 150, 70
    gap = 30
    total = len(steps) * box_w + (len(steps) - 1) * gap
    x = (w - total) // 2
    y = 50
    for i, step in enumerate(steps):
        fill = "#1E73E8" if i in (0, 5) else "#E8F1FC"
        tc = "#FFFFFF" if i in (0, 5) else "#0F4CA8"
        draw_rounded_rect(draw, (x, y, x + box_w, y + box_h), 10, fill, "#0F4CA8", 2)
        lines = step.split("\n")
        lh = 18
        ty = y + (box_h - lh * len(lines)) // 2
        for li, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
            draw.text((x + (box_w - tw) // 2, ty + li * lh), line, fill=tc, font=font)
        if i < len(steps) - 1:
            ax0 = x + box_w + 4
            ax1 = x + box_w + gap - 4
            mid = y + box_h // 2
            draw.line((ax0, mid, ax1 - 8, mid), fill="#1E73E8", width=3)
            draw.polygon(
                [(ax1 - 10, mid - 7), (ax1 - 10, mid + 7), (ax1, mid)],
                fill="#1E73E8",
            )
        x += box_w + gap
    img.save(path, "PNG")


def make_era_pipeline_image(path: Path):
    w, h = 1100, 420
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    font = _font(15, bold=True)
    font_sm = _font(11)
    stages = [
        ("1. Retrieval", "835 ERA · Portal EOBs\nPaper scan · Bank feed"),
        ("2. Extraction", "Normalize remittance\nCARC/RARC · Confidence"),
        ("3. Matching", "Claim match · Trace #\nThree-way reconcile"),
        ("4. Rules Engine", "Posting policy\nAuto · Flag · Review"),
        ("5. OD Posting", "ClaimProcs · Payments\nWrite-offs · Notes"),
        ("6. Exceptions", "Aging queues\nHuman ownership"),
    ]
    cols = 3
    box_w, box_h = 300, 88
    margin_x = 60
    gap_x = 40
    gap_y = 28
    start_y = 24
    for i, (title, body) in enumerate(stages):
        col = i % cols
        row = i // cols
        x = margin_x + col * (box_w + gap_x)
        y = start_y + row * (box_h + gap_y + 20)
        draw_rounded_rect(draw, (x, y, x + box_w, y + box_h), 10, "#E8F1FC", "#1E73E8", 2)
        draw.rectangle((x, y, x + 10, y + box_h), fill="#1E73E8")
        draw.text((x + 24, y + 12), title, fill="#0F4CA8", font=font)
        draw.multiline_text((x + 24, y + 38), body, fill="#5A6A7A", font=font_sm, spacing=3)
    footer = "Planned ezFi remittance automation — structured first, unstructured with confidence gates"
    bbox = draw.textbbox((0, 0), footer, font=font_sm)
    tw = bbox[2] - bbox[0]
    draw.text(((w - tw) // 2, h - 28), footer, fill="#5A6A7A", font=font_sm)
    img.save(path, "PNG")


# ── Section writers ──────────────────────────────────────────────────────────
def configure_section(section, *, first_page_different=True):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.7)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = first_page_different

    # Header (subsequent pages)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hp, before=0, after=0)
    run = hp.add_run()
    run.add_picture(str(LOGO), width=Inches(0.72))
    hp2 = header.add_paragraph()
    hp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(hp2, before=0, after=1)
    r2 = hp2.add_run("AI-Powered Insurance Verification & RCM Platform")
    set_run_font(r2, size=8, color=GRAY)
    add_horizontal_line(hp2, "1E73E8", "8")

    fh = section.first_page_header
    fh.is_linked_to_previous = False
    fh.paragraphs[0].text = ""

    for footer in (section.footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.text = ""
        table = footer.add_table(1, 3, Inches(7.0))
        table.autofit = True
        remove_table_borders(table)
        left, mid, right = table.rows[0].cells
        pl = left.paragraphs[0]
        set_paragraph_spacing(pl, before=2, after=0)
        rl = pl.add_run("ezFi  ·  Confidential")
        set_run_font(rl, size=8, color=GRAY)
        pm = mid.paragraphs[0]
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(pm, before=2, after=0)
        rm = pm.add_run(FOOTER_TEXT)
        set_run_font(rm, size=8, bold=True, color=BLUE)
        pr = right.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(pr, before=2, after=0)
        rr = pr.add_run("Page ")
        set_run_font(rr, size=8, color=GRAY)
        add_page_number(pr)
        for run in pr.runs[1:]:
            set_run_font(run, size=8, color=GRAY)


def build_cover(doc):
    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)
        p.add_run("")

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(logo_p, before=36, after=8)
    logo_p.add_run().add_picture(str(LOGO), width=Inches(2.5))

    # Short centered accent via a 3-column spacer table
    accent = doc.add_table(rows=1, cols=3)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(accent)
    remove_table_borders(accent)
    accent.rows[0].cells[0].width = Inches(2.45)
    mid = accent.rows[0].cells[1]
    mid.width = Inches(1.8)
    shade_cell(mid, HEADER_BG)
    set_cell_borders(mid, HEADER_BG, "2")
    mp = mid.paragraphs[0]
    set_paragraph_spacing(mp, before=1, after=1)
    mr = mp.add_run(" ")
    set_run_font(mr, size=1)
    accent.rows[0].cells[2].width = Inches(2.45)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=12, after=4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=8, after=8)
    tr = title.add_run("AI-Powered Insurance Verification\n& Revenue Cycle Automation Platform")
    set_run_font(tr, size=20, bold=True, color=BLUE_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, before=4, after=10)
    sr = subtitle.add_run("Enterprise Automation for Modern Dental Organizations")
    set_run_font(sr, size=13, color=GRAY)

    tag = doc.add_table(rows=1, cols=1)
    tag.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc = tag.cell(0, 0)
    shade_cell(tc, LIGHT_BG)
    set_cell_borders(tc, "D0D7DE", "4")
    tp = tc.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(tp, before=6, after=6)
    tr2 = tp.add_run(
        "Eligibility Verification  ·  OpenDental Writeback  ·  Voice AI  ·  RCM Roadmap"
    )
    set_run_font(tr2, size=9.5, color=BLUE_MID)

    for _ in range(4):
        sp = doc.add_paragraph()
        set_paragraph_spacing(sp, before=0, after=0)

    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(audience, before=0, after=4)
    ar = audience.add_run(
        "Prepared for DSOs  ·  Enterprise Dental Groups  ·  Practice Owners\n"
        "Investors  ·  Strategic Technology Partners"
    )
    set_run_font(ar, size=10, color=GRAY)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(foot, before=16, after=0)
    fr = foot.add_run(FOOTER_TEXT)
    set_run_font(fr, size=11, bold=True, color=BLUE)

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(conf, before=4, after=0)
    cr = conf.add_run("Executive Product Overview  ·  Confidential")
    set_run_font(cr, size=9, color=GRAY)

    doc.add_page_break()


def build_toc(doc):
    add_heading_styled(doc, "Contents", 1)
    items = [
        ("01", "Executive Summary"),
        ("02", "Why Insurance Verification Matters"),
        ("03", "ezFi Platform Overview"),
        ("04", "Eligibility Verification"),
        ("05", "Specialist-Level Verification of Benefits"),
        ("06", "OpenDental Writeback"),
        ("07", "Safety & Data Integrity"),
        ("08", "Voice Verification"),
        ("09", "Future Revenue Cycle Roadmap"),
        ("10", "EOB / ERA Automation"),
        ("11", "Product Roadmap"),
        ("12", "Closing"),
    ]
    table = doc.add_table(rows=len(items), cols=2)
    set_table_fixed(table)
    remove_table_borders(table)
    for i, (num, title) in enumerate(items):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.width = Inches(0.55)
        c1.width = Inches(6.0)
        if i % 2 == 0:
            shade_cell(c0, ACCENT_BG)
            shade_cell(c1, ACCENT_BG)
        p0 = c0.paragraphs[0]
        set_paragraph_spacing(p0, before=3, after=3)
        r0 = p0.add_run(num)
        set_run_font(r0, size=10, bold=True, color=BLUE)
        p1 = c1.paragraphs[0]
        set_paragraph_spacing(p1, before=3, after=3)
        r1 = p1.add_run(title)
        set_run_font(r1, size=10.5, color=BLACK)
    doc.add_page_break()


def section_exec_summary(doc):
    add_heading_styled(doc, "1. Executive Summary", 1)
    add_body(
        doc,
        "Dental insurance verification remains one of the largest administrative burdens "
        "in modern practice operations. Front-desk and billing teams spend hours each day "
        "calling payers, interpreting benefits, updating OpenDental, and explaining patient "
        "responsibility—work that is slow, inconsistent, and difficult to scale across locations.",
        after=6,
    )
    add_body(
        doc,
        "ezFi is an AI-powered platform that automates insurance verification, structured "
        "documentation, and OpenDental writeback—while keeping clinical and administrative "
        "staff in control of every final decision.",
        after=8,
    )
    add_kpi_row(
        doc,
        [
            ("Flagship Capability", "Eligibility\nVerification"),
            ("Native Integration", "OpenDental\nWriteback"),
            ("Human Control", "Review &\nApproval"),
            ("Roadmap", "End-to-End\nRCM"),
        ],
    )
    add_callout(
        doc,
        "Platform Philosophy",
        "Automate the work. Preserve clinical and financial judgment. Never overwrite what "
        "staff have carefully entered. OpenDental remains the system of record; ezFi is the "
        "intelligence and orchestration layer that makes verification accurate, auditable, and fast.",
    )
    add_body(doc, "With ezFi, organizations can:", after=2)
    for item in (
        "Reduce phone-time and manual benefit interpretation",
        "Produce specialist-quality Verification of Benefits summaries from electronic eligibility",
        "Write structured results safely into OpenDental",
        "Extend incomplete electronic responses with Voice AI",
        "Build toward claim preparation, ERA/EOB posting, and denial management on the same platform",
    ):
        add_bullet(doc, item)
    # Continue into next section on same page if room; break for visual chaptering
    doc.add_page_break()


def section_why(doc, flow_img: Path):
    add_heading_styled(doc, "2. Why Insurance Verification Matters", 1)
    add_body(
        doc,
        "Accurate insurance verification sits at the front of every successful revenue cycle. "
        "When benefits are wrong or incomplete, estimates fail, collections suffer, and patients "
        "lose trust at the chairside conversation that matters most.",
        after=6,
    )
    add_heading_styled(doc, "Traditional Workflow", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=4)
    p.add_run().add_picture(str(flow_img), width=Inches(3.6))

    add_heading_styled(doc, "Where the Traditional Model Breaks Down", 2)
    add_styled_table(
        doc,
        ["Challenge", "Operational Impact"],
        [
            ["Manual payer calls", "Hours of hold time; inconsistent answers across staff"],
            ["Benefit interpretation", "Variable quality; tribal knowledge; training burden"],
            ["OpenDental updates", "Incomplete notes; missed fields; estimate drift"],
            ["Patient communication", "Surprise balances; delayed treatment acceptance"],
            ["Multi-location scale", "No shared standard; hard to audit or measure"],
        ],
        col_widths=[Inches(2.2), Inches(4.5)],
    )
    add_callout(
        doc,
        "How ezFi Modernizes Verification",
        "ezFi replaces fragmented phone-and-notes workflows with a layered eligibility engine, "
        "structured Verification of Benefits output, confidence scoring, and controlled OpenDental "
        "writeback—so every location verifies the same way, with the same documentation standard.",
        spacer=False,
    )
    doc.add_page_break()


def section_platform(doc):
    add_heading_styled(doc, "3. ezFi Platform Overview", 1)
    add_body(
        doc,
        "ezFi sits between OpenDental and the payer ecosystem as an AI orchestration layer. "
        "Eligibility Verification is the flagship capability today. Voice Verification and "
        "OpenDental Writeback complete the live operating stack. Additional revenue cycle agents "
        "extend the same architecture over time.",
        after=6,
    )
    add_architecture_diagram(doc)
    add_heading_styled(doc, "Design Principles", 2)
    add_styled_table(
        doc,
        ["Principle", "Meaning"],
        [
            ["OpenDental as system of record", "ezFi reads and writes; it does not replace the PMS"],
            ["Human-in-the-loop by default", "Staff approve consequential actions"],
            ["Confidence before mutation", "Low-confidence results route to review"],
            ["Preserve manual edits", "Staff corrections are never blindly overwritten"],
            ["Auditability", "Every verification and writeback is traceable"],
        ],
        col_widths=[Inches(2.6), Inches(4.1)],
    )
    doc.add_page_break()


def section_eligibility(doc):
    add_heading_styled(doc, "4. Eligibility Verification", 1)
    add_body(
        doc,
        "Eligibility Verification is ezFi’s flagship capability—a complete, layered pipeline "
        "that turns a patient and plan into a specialist-quality Verification of Benefits, "
        "ready for staff review and safe OpenDental synchronization.",
        after=6,
    )
    add_heading_styled(doc, "Layered Architecture", 2)
    add_layer_diagram(doc)

    add_heading_styled(doc, "Pipeline Detail", 2)
    add_styled_table(
        doc,
        ["Layer", "What ezFi Does"],
        [
            [
                "Patient validation",
                "Confirms identity and demographics against OpenDental before payer inquiry",
            ],
            [
                "Insurance validation",
                "Validates plan linkage, subscriber relationships, and coverage order",
            ],
            [
                "Carrier validation",
                "Resolves payer identity and routing for electronic eligibility",
            ],
            [
                "OpenDental integration",
                "Pulls live plan context and prepares deterministic writeback targets",
            ],
            [
                "Stedi 270/271 eligibility",
                "Issues electronic benefit inquiry and captures structured 271 responses",
            ],
            [
                "Eligibility normalization",
                "Maps heterogeneous payer responses into a canonical dental benefits model",
            ],
            [
                "Coverage interpretation",
                "Derives percentages, accumulators, limitations, and plan clauses",
            ],
            [
                "Business rules",
                "Applies organization and payer-specific policies before staff see results",
            ],
            [
                "Caching",
                "Reuses recent, still-valid eligibility where appropriate to reduce cost and latency",
            ],
            [
                "Confidence scoring",
                "Scores completeness and reliability; gates automated vs. review paths",
            ],
            [
                "Review queue",
                "Surfaces incomplete, conflicting, or low-confidence cases for staff",
            ],
            [
                "Dashboard",
                "Operational workspace for queue management, detail review, and approval",
            ],
            [
                "Exception handling",
                "Routes failures, payer gaps, and writeback conflicts into owned exception queues",
            ],
        ],
        col_widths=[Inches(2.2), Inches(4.5)],
    )
    add_callout(
        doc,
        "Outcome",
        "Staff receive a consistent, specialist-level Verification of Benefits—not a raw 271 dump—"
        "with clear confidence signals and a controlled path into OpenDental.",
        spacer=False,
    )
    doc.add_page_break()


def section_vob(doc):
    add_heading_styled(doc, "5. Specialist-Level Verification of Benefits", 1)
    add_body(
        doc,
        "Raw electronic eligibility is incomplete for dental operations. ezFi transforms 271 data "
        "into structured, specialist-quality Verification of Benefits summaries that front desk, "
        "treatment coordinators, and billing teams can act on immediately.",
        after=6,
    )
    add_heading_styled(doc, "Extracted & Interpreted Benefit Elements", 2)
    add_benefit_grid(doc)
    add_heading_styled(doc, "From Raw 271 to Operational Clarity", 2)
    add_styled_table(
        doc,
        ["Source Reality", "ezFi Output"],
        [
            ["Fragmented EDI segments", "Canonical dental benefits record"],
            ["Payer-specific naming", "Normalized coverage categories"],
            ["Missing narrative context", "Deterministic VOB summary for Benefit Notes"],
            ["Unclear patient cost", "Procedure-aware patient estimates where data allows"],
            ["No audit trail", "Verification timestamp, check ID, and review status"],
        ],
        col_widths=[Inches(2.6), Inches(4.1)],
    )
    add_callout(
        doc,
        "Specialist Standard",
        "ezFi is designed to produce the same depth a skilled insurance coordinator would "
        "document after a thorough payer call—coverage, deductibles, maximums, limitations, "
        "clauses, and patient estimates—then keep that standard consistent across every location.",
        spacer=False,
    )
    doc.add_page_break()


def section_writeback(doc):
    add_heading_styled(doc, "6. OpenDental Writeback", 1)
    add_body(
        doc,
        "Verification only creates value when it lands where staff already work. ezFi writes "
        "structured eligibility results into OpenDental across multiple, purpose-built locations—"
        "each with clear intent, isolation of failures, and safeguards for manually curated data.",
        after=6,
    )
    add_heading_styled(doc, "Writeback Locations", 2)
    add_styled_table(
        doc,
        ["OpenDental Location", "What ezFi Writes", "Purpose"],
        [
            [
                "Benefit Notes",
                "Complete deterministic VOB summary: coverage, deductibles, annual maximums, "
                "patient estimates, verification timestamp",
                "Primary structured snapshot for clinical and billing review",
            ],
            [
                "Subscriber Notes",
                "Short insurance summary visible on the insurance grid",
                "At-a-glance status for front desk",
            ],
            [
                "Insurance Verification",
                "Verification status update (Patient Enrollment & Insurance Benefit)",
                "Audit trail of when and how benefits were verified",
            ],
            [
                "Communication Log",
                "Front desk documentation and estimated patient responsibility",
                "Patient-facing conversation support",
            ],
            [
                "Benefits Grid",
                "Coverage percentages, deductibles, annual maximums; preserves manually edited rows",
                "Structured plan benefits for estimates and claims",
            ],
            [
                "Insurance Adjustments",
                "Remaining vs. used totals when sufficient payer information exists",
                "Financial accumulator sync—only with adequate data",
            ],
        ],
        col_widths=[Inches(1.7), Inches(2.9), Inches(2.1)],
    )
    add_heading_styled(doc, "Writeback Execution Model", 2)
    add_body(
        doc,
        "Each writeback step is isolated. A failure in one location does not abort the others. "
        "Financial and benefits-grid mutations support dry-run / shadow comparison and "
        "confidence gating before live mutation.",
        after=6,
    )
    add_callout(
        doc,
        "OpenDental Remains Authoritative",
        "ezFi never invents fee schedules, ledger payments, or fabricated adjustments. "
        "Writeback is deterministic, reversible where snapshots exist, and designed for "
        "enterprise trust—not silent ledger mutation.",
        spacer=False,
    )
    doc.add_page_break()


def section_safety(doc):
    add_heading_styled(doc, "7. Safety & Data Integrity", 1)
    add_body(
        doc,
        "Automation without control is a liability. ezFi is engineered so clinics retain "
        "complete authority over financial and benefits data while still capturing the speed "
        "and consistency of AI-assisted verification.",
        after=6,
    )
    add_heading_styled(doc, "ezFi Intentionally Does Not", 2)
    add_styled_table(
        doc,
        ["Prohibited Action", "Why It Matters"],
        [
            ["Overwrite manually edited benefits", "Staff corrections reflect real-world plan knowledge"],
            ["Fabricate adjustments", "Ledger integrity must never be invented by a model"],
            ["Fabricate payments", "Cash and remittance must originate from real payer activity"],
            ["Create fake ledger transactions", "OpenDental financials stay audit-safe"],
            ["Auto-assign fee schedules", "Contractual fees remain a human / contract decision"],
        ],
        col_widths=[Inches(2.8), Inches(3.9)],
    )
    add_heading_styled(doc, "Control Surfaces", 2)
    controls = [
        (
            "Shadow Mode",
            "Run full verification and proposed writeback without mutating OpenDental. Compare before go-live.",
        ),
        (
            "Review Queue",
            "Low-confidence, incomplete, or policy-sensitive results wait for staff judgment.",
        ),
        (
            "Exception Queue",
            "Failures and conflicts are aged, owned, and actionable—never lost in inboxes.",
        ),
        (
            "Approval Workflow",
            "Consequential writebacks proceed only after configured review, preserving control.",
        ),
    ]
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    for i, (title, body) in enumerate(controls):
        cell = table.cell(i // 2, i % 2)
        shade_cell(cell, CALLOUT_BG)
        set_cell_borders(cell, "1E73E8", "8")
        cell.width = Inches(3.35)
        p1 = cell.paragraphs[0]
        set_paragraph_spacing(p1, before=4, after=1)
        r1 = p1.add_run(title)
        set_run_font(r1, size=10, bold=True, color=BLUE_DARK)
        p2 = cell.add_paragraph()
        set_paragraph_spacing(p2, before=0, after=4)
        r2 = p2.add_run(body)
        set_run_font(r2, size=9, color=GRAY)
    # No forced page break — continue into Voice on same/next natural page


def section_voice(doc, voice_img: Path):
    add_heading_styled(doc, "8. Voice Verification", 1)
    add_body(
        doc,
        "Electronic eligibility is powerful—but not always complete. Voice AI extends the "
        "eligibility stack when the 271 leaves gaps that used to require a staff phone call.",
        after=6,
    )
    add_heading_styled(doc, "Workflow", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=6)
    p.add_run().add_picture(str(voice_img), width=Inches(6.5))

    add_heading_styled(doc, "Conversation Design", 2)
    add_styled_table(
        doc,
        ["Capability", "Description"],
        [
            [
                "Context-aware conversations",
                "The agent knows the patient, plan, and missing benefit fields before dialing",
            ],
            [
                "One question at a time",
                "Payer agents receive clear, sequential questions—not overwhelming scripts",
            ],
            [
                "Intelligent follow-up",
                "Answers drive the next question; ambiguity is clarified, not ignored",
            ],
            [
                "Dashboard review",
                "Transcripts and extracted fields land in the same review workspace as electronic checks",
            ],
            [
                "Safe reconciliation",
                "Voice findings merge into the VOB and writeback path only after review and approval",
            ],
        ],
        col_widths=[Inches(2.3), Inches(4.4)],
    )
    add_callout(
        doc,
        "Complement, Don’t Replace",
        "Voice Verification is an extension of electronic eligibility—not a parallel process. "
        "The destination is the same: a complete VOB, staff control, and safe OpenDental writeback.",
        spacer=False,
    )
    doc.add_page_break()


def section_future_rcm(doc):
    add_heading_styled(doc, "9. Future Revenue Cycle Roadmap", 1)
    add_body(
        doc,
        "Eligibility Verification is the foundation. The same architecture—OpenDental as "
        "system of record, confidence gating, human review, and deterministic writeback—"
        "extends across the dental revenue cycle.",
        after=6,
    )
    # Three columns for coding / prior auth / claim prep
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    columns = [
        (
            "Coding Agent",
            [
                "AI-assisted CDT coding",
                "Documentation review",
                "Clinical narratives",
            ],
        ),
        (
            "Prior Authorization",
            [
                "Auth requirement detection",
                "Attachments & submission",
                "Tracking & renewals",
            ],
        ),
        (
            "Claim Preparation",
            [
                "Pre-submission validation",
                "Missing documentation",
                "Readiness scoring",
            ],
        ),
    ]
    for i, (title, items) in enumerate(columns):
        cell = table.cell(0, i)
        shade_cell(cell, LIGHT_BG)
        set_cell_borders(cell, "1E73E8", "8")
        cell.width = Inches(2.25)
        ph = cell.paragraphs[0]
        set_paragraph_spacing(ph, before=4, after=3)
        rh = ph.add_run(title)
        set_run_font(rh, size=10, bold=True, color=BLUE_DARK)
        for item in items:
            px = cell.add_paragraph()
            set_paragraph_spacing(px, before=1, after=1)
            rx = px.add_run(f"•  {item}")
            set_run_font(rx, size=9, color=BLACK)
        sp = cell.add_paragraph()
        set_paragraph_spacing(sp, before=0, after=3)
    add_callout(
        doc,
        "One Operating Model",
        "Each future agent reuses ezFi’s review queues, confidence thresholds, audit trail, "
        "and OpenDental synchronization patterns—so DSOs do not adopt a patchwork of point tools.",
        spacer=False,
    )
    # Continue into ERA section


def section_era(doc, era_img: Path):
    add_heading_styled(doc, "10. EOB / ERA Automation", 1)
    add_body(
        doc,
        "As ezFi expands from verification into full revenue cycle automation, remittance "
        "processing becomes a natural next layer—structured data first, confidence gates, "
        "deterministic OpenDental posting, and human ownership of exceptions.",
        after=6,
    )
    add_heading_styled(doc, "Planned Remittance Pipeline", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=6)
    p.add_run().add_picture(str(era_img), width=Inches(6.3))

    add_heading_styled(doc, "Capability Overview", 2)
    add_styled_table(
        doc,
        ["Capability", "ezFi Approach"],
        [
            [
                "ERA ingestion",
                "Prefer X12 835 / ERA as the highest-fidelity input; maximize electronic remittance share",
            ],
            [
                "EOB interpretation",
                "Document understanding for paper and portal EOBs, with field-level confidence scores",
            ],
            [
                "Payment posting",
                "Deterministic OpenDental API posting to claimprocs—never generative ledger writes",
            ],
            [
                "Payment reconciliation",
                "Three-way match: remittance total = ledger payment = bank credit (EFT trace / check #)",
            ],
            [
                "Exception handling",
                "Named, aging queues for unmatched remits, denials, downgrades, takebacks, and refunds",
            ],
            [
                "Human review",
                "Auto-post only when confidence, match, and rules thresholds are met; otherwise review",
            ],
            [
                "OpenDental sync",
                "Payments, write-offs, deductibles, source documents, and posting rationale notes",
            ],
        ],
        col_widths=[Inches(2.0), Inches(4.7)],
    )
    add_heading_styled(doc, "Posting Integrity Principles", 2)
    for item in (
        "Rules engine primary for write-off vs. patient responsibility vs. denial classification",
        "Procedure-level posting for PPO accuracy—not lump totals that corrupt A/R",
        "No silent netting of takebacks; reverse with reference to the original claim",
        "Immutable audit: source document, rule version, confidence, before/after ledger state",
    ):
        add_bullet(doc, item)
    add_callout(
        doc,
        "Part of the Same Platform",
        "EOB/ERA automation is the remittance chapter of ezFi’s end-to-end revenue cycle platform—"
        "built on the trust model proven in Eligibility Verification and OpenDental Writeback.",
        spacer=False,
    )
    doc.add_page_break()


def section_roadmap(doc):
    add_heading_styled(doc, "11. Product Roadmap", 1)
    add_body(
        doc,
        "ezFi delivers value now in verification and writeback, then expands along a clear "
        "path toward enterprise DSO revenue cycle operations.",
        after=8,
    )
    add_timeline(doc)
    add_heading_styled(doc, "Investment Thesis for Enterprise Buyers", 2)
    add_styled_table(
        doc,
        ["Horizon", "Buyer Outcome"],
        [
            ["Current", "Standardize verification quality; cut admin time; harden OpenDental documentation"],
            ["Next", "Reduce coding and auth friction; improve claim readiness before submission"],
            ["Future", "Compress remittance-to-cash; control denials; measure RCM performance at scale"],
        ],
        col_widths=[Inches(1.4), Inches(5.3)],
    )
    # Continue into closing


def section_closing(doc):
    add_heading_styled(doc, "12. Closing", 1)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, ACCENT_BG)
    set_cell_borders(cell, HEADER_BG, "18")
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=10, after=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "ezFi combines intelligent insurance verification, OpenDental integration, "
        "and AI-driven automation to reduce administrative burden while helping dental "
        "organizations improve efficiency, accuracy, and patient financial transparency."
    )
    set_run_font(r, size=12, bold=False, color=BLUE_DARK)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=10, after=4)

    add_heading_styled(doc, "Engage With ezFi", 2)
    add_body(
        doc,
        "Whether you operate a single high-volume specialty practice or a multi-state DSO, "
        "ezFi is built to meet you where you work today—in OpenDental—and to grow with your "
        "revenue cycle ambitions.",
        after=6,
    )
    add_kpi_row(
        doc,
        [
            ("Start With", "Eligibility\nVerification"),
            ("Integrate Via", "OpenDental\nWriteback"),
            ("Scale With", "Voice +\nRCM Agents"),
            ("Govern With", "Review &\nShadow Mode"),
        ],
    )

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(logo_p, before=12, after=6)
    logo_p.add_run().add_picture(str(LOGO), width=Inches(1.4))

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(foot, before=4, after=2)
    fr = foot.add_run(FOOTER_TEXT)
    set_run_font(fr, size=11, bold=True, color=BLUE)

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(conf, before=2, after=0)
    cr = conf.add_run("© Smile Suites  ·  ezFi Executive Product Overview  ·  Confidential")
    set_run_font(cr, size=9, color=GRAY)


def main():
    ASSETS.mkdir(parents=True, exist_ok=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)

    if not LOGO.exists():
        raise SystemExit(f"Logo not found: {LOGO}")

    flow_img = ASSETS / "traditional_flow.png"
    voice_img = ASSETS / "voice_flow.png"
    era_img = ASSETS / "era_pipeline.png"
    make_traditional_flow_image(flow_img)
    make_voice_flow_image(voice_img)
    make_era_pipeline_image(era_img)

    doc = Document()
    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    section = doc.sections[0]
    configure_section(section, first_page_different=True)

    build_cover(doc)
    build_toc(doc)
    section_exec_summary(doc)
    section_why(doc, flow_img)
    section_platform(doc)
    section_eligibility(doc)
    section_vob(doc)
    section_writeback(doc)
    section_safety(doc)
    section_voice(doc, voice_img)
    section_future_rcm(doc)
    section_era(doc, era_img)
    section_roadmap(doc)
    section_closing(doc)

    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
